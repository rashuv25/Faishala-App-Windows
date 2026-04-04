# -*- coding: utf-8 -*-
"""DOCX document generation."""

from pathlib import Path
from typing import Dict, Any

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config.settings import Settings
from config.constants import (
    DOCUMENT_TOP_BLANK_LINES,
    NEPALI_MONTHS,
    DEFAULT_TAPSIL_POINTS
)
from utils.nepali_text import sanitize_document_data, sanitize_nepali_text


class DocxGenerator:
    """Generates DOCX documents."""

    def __init__(self):
        """Initialize DOCX generator."""
        self.font_name = Settings.EXPORT_FONT_NAME

    def generate(self, document_data: Dict[str, Any], output_path: Path) -> bool:
        """Generate DOCX file from document data."""
        try:
            doc = Document()
            self._set_page_layout(doc)
            cleaned_data = sanitize_document_data(document_data)

            self._set_default_font(doc)

            for _ in range(DOCUMENT_TOP_BLANK_LINES):
                doc.add_paragraph()

            self._add_header(doc, cleaned_data)
            self._add_wadi_pratiwadi(doc, cleaned_data)
            self._add_mudda_section(doc, cleaned_data)
            self._add_case_points(doc, cleaned_data)
            self._add_office_decision(doc, cleaned_data)
            self._add_tapsil(doc, cleaned_data)
            self._add_footer(doc, cleaned_data)

            doc.save(str(output_path))
            return True

        except Exception as e:
            print(f"DOCX generation error: {e}")
            return False

    def _set_page_layout(self, doc: Document) -> None:
        """Set document page layout."""
        for section in doc.sections:
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(0.8)
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)

    def _set_default_font(self, doc: Document) -> None:
        """Set default font for document."""
        style = doc.styles["Normal"]
        font = style.font
        font.name = self.font_name
        font.size = Pt(14)

        r = style.element
        rPr = r.get_or_add_rPr()

        for child in list(rPr):
            if child.tag == qn("w:rFonts"):
                rPr.remove(child)

        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), self.font_name)
        rFonts.set(qn("w:hAnsi"), self.font_name)
        rFonts.set(qn("w:cs"), self.font_name)
        rFonts.set(qn("w:eastAsia"), self.font_name)
        rPr.append(rFonts)

    def _apply_font_to_run(self, run, bold: bool = False, size: int = 14) -> None:
        """Apply export font to an individual run."""
        run.font.name = self.font_name
        run.font.size = Pt(size)
        run.bold = bold

        r = run._element
        rPr = r.get_or_add_rPr()

        for child in list(rPr):
            if child.tag == qn("w:rFonts"):
                rPr.remove(child)

        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), self.font_name)
        rFonts.set(qn("w:hAnsi"), self.font_name)
        rFonts.set(qn("w:cs"), self.font_name)
        rFonts.set(qn("w:eastAsia"), self.font_name)
        rPr.append(rFonts)

    def _add_text_run(self, paragraph, text: str, bold: bool = False, size: int = 14):
        """Add a sanitized text run to a paragraph."""
        run = paragraph.add_run(sanitize_nepali_text(str(text)))
        self._apply_font_to_run(run, bold=bold, size=size)
        return run

    def _add_text_with_dot_leader(self, paragraph, left_text: str, bold: bool = True, size: int = 14):
        """Add text followed by a dynamic right-side dotted leader."""
        tabs = paragraph.paragraph_format.tab_stops
        tabs.add_tab_stop(
            Cm(18.2),
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.DOTS
        )
        self._add_text_run(paragraph, left_text, bold=bold, size=size)
        self._add_text_run(paragraph, "\t", bold=False, size=size)

    def _add_tapsil_line(self, paragraph, text: str, point_num: str, size: int = 14):
        """Add a tapsil line with dynamic dotted leader and far-right point number."""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.left_indent = Cm(1.0)
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Cm(18.2),
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.DOTS
        )

        self._add_text_run(paragraph, text, size=size)
        self._add_text_run(paragraph, "\t", size=size)
        self._add_text_run(paragraph, point_num, size=size)

    def _set_cell_margins(self, cell, top=80, start=80, bottom=80, end=80):
        """Set table cell margins."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        tcMar = None
        for child in tcPr:
            if child.tag == qn("w:tcMar"):
                tcMar = child
                break

        if tcMar is None:
            tcMar = OxmlElement("w:tcMar")
            tcPr.append(tcMar)

        for m, val in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
            node = None
            for child in tcMar:
                if child.tag == qn(f"w:{m}"):
                    node = child
                    break
            if node is None:
                node = OxmlElement(f"w:{m}")
                tcMar.append(node)
            node.set(qn("w:w"), str(val))
            node.set(qn("w:type"), "dxa")

    def _add_header(self, doc: Document, data: Dict) -> None:
        """Add document header."""
        district = data.get("district", "")
        cdo_name = data.get("cdo_name", "")

        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p1.paragraph_format.left_indent = Cm(1.0)
        p1.paragraph_format.right_indent = Cm(0)
        p1.paragraph_format.space_after = Pt(4)
        p1.paragraph_format.line_spacing = 1.0

        self._add_text_run(p1, "जिल्ला प्रशासन कार्यालय, ", bold=False, size=13)
        self._add_text_run(p1, district, bold=False, size=13)
        self._add_text_run(p1, "का प्रमुख जिल्ला अधिकारी ", bold=False, size=13)
        self._add_text_run(p1, cdo_name, bold=True, size=13)
        self._add_text_run(p1, " को", bold=False, size=13)

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p2.paragraph_format.left_indent = Cm(1.0)
        p2.paragraph_format.right_indent = Cm(0)
        p2.paragraph_format.space_after = Pt(10)
        p2.paragraph_format.line_spacing = 1.0

        self._add_text_with_dot_leader(p2, "इजालासबाट भएको फैसला", bold=True, size=13)

    def _add_wadi_pratiwadi(self, doc: Document, data: Dict) -> None:
        """Add Wadi and Pratiwadi section."""
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False

        left_cell = table.cell(0, 0)
        right_cell = table.cell(0, 1)

        left_cell.width = Cm(7.0)
        right_cell.width = Cm(10.4)

        self._set_cell_margins(left_cell, start=20, end=10)
        self._set_cell_margins(right_cell, start=220, end=10)

        # Left side
        left_cell.text = ""
        lp0 = left_cell.paragraphs[0]
        lp0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp0.paragraph_format.space_after = Pt(6)
        self._add_text_run(lp0, "वादी", bold=True, size=15)

        for line in str(data.get("wadi_content", "")).split("\n"):
            line = line.strip()
            if not line:
                continue
            lp = left_cell.add_paragraph()
            lp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            lp.paragraph_format.left_indent = Cm(0)
            lp.paragraph_format.right_indent = Cm(0)
            lp.paragraph_format.space_after = Pt(2)
            self._add_text_run(lp, line, size=14)

        # Right side
        right_cell.text = ""
        rp0 = right_cell.paragraphs[0]
        rp0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rp0.paragraph_format.space_after = Pt(6)
        self._add_text_run(rp0, "प्रतिवादी", bold=True, size=15)

        for line in str(data.get("pratiwadi_content", "")).split("\n"):
            line = line.strip()
            if not line:
                continue
            rp = right_cell.add_paragraph()
            rp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            rp.paragraph_format.left_indent = Cm(0.8)
            rp.paragraph_format.right_indent = Cm(0)
            rp.paragraph_format.space_after = Pt(2)
            self._add_text_run(rp, line, size=14)

    def _add_mudda_section(self, doc: Document, data: Dict) -> None:
        """Add Mudda section."""
        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_after = Pt(8)
        self._add_text_run(p1, "मुद्दा : ", bold=True, size=14)
        self._add_text_run(p1, data.get("mudda", ""), size=14)

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(12)
        self._add_text_run(p2, "मु. द. नं. : ", bold=True, size=14)
        self._add_text_run(p2, data.get("mudda_number", ""), size=14)

    def _add_case_points(self, doc: Document, data: Dict) -> None:
        """Add case points section."""
        title = doc.add_paragraph()
        self._add_text_run(title, "केसको संक्षिप्त व्यहोरा", bold=True, size=15)

        case_points = data.get("case_points", [""])
        for i, point in enumerate(case_points, 1):
            nepali_num = self._to_nepali_number(i)

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Cm(1.2)
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.tab_stops.add_tab_stop(Cm(2.3))

            self._add_text_run(p, f"{nepali_num}.", bold=False, size=14)
            self._add_text_run(p, "\t", size=14)
            self._add_text_run(p, point, size=14)

    def _add_office_decision(self, doc: Document, data: Dict) -> None:
        """Add office decision section."""
        title = doc.add_paragraph()
        self._add_text_run(title, "कार्यालयको ठहर", bold=True, size=15)

        decision = data.get("office_decision", "")
        for para in str(decision).split("\n"):
            para = para.strip()
            if not para:
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.first_line_indent = Cm(1.2)
            p.paragraph_format.space_after = Pt(4)
            self._add_text_run(p, para, size=14)

    def _add_tapsil(self, doc: Document, data: Dict) -> None:
        """Add tapsil section."""
        title = doc.add_paragraph()
        title.paragraph_format.space_after = Pt(8)
        self._add_text_run(title, "तपसिल", bold=True, size=15)

        tapsil_points = data.get("tapsil_points", DEFAULT_TAPSIL_POINTS)
        for i, point in enumerate(tapsil_points, 1):
            nepali_num = self._to_nepali_number(i)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            self._add_tapsil_line(p, sanitize_nepali_text(point), nepali_num, size=14)

        gap1 = doc.add_paragraph()
        gap1.paragraph_format.space_after = Pt(0)
        gap2 = doc.add_paragraph()
        gap2.paragraph_format.space_after = Pt(0)

    def _add_footer(self, doc: Document, data: Dict) -> None:
        """Add document footer."""
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False

        left_cell = table.cell(0, 0)
        right_cell = table.cell(0, 1)

        left_cell.width = Cm(8.0)
        right_cell.width = Cm(9.4)

        self._set_cell_margins(left_cell)
        self._set_cell_margins(right_cell)

        left_cell.text = ""
        p_left = left_cell.paragraphs[0]
        p_left.paragraph_format.space_after = Pt(0)
        self._add_text_run(
            p_left,
            f"टिपोट गर्ने ना.सु. {data.get('typist_name', '')}",
            size=14
        )

        right_cell.text = ""

        p1 = right_cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p1.paragraph_format.space_after = Pt(0)
        p1.paragraph_format.line_spacing = 1.0
        self._add_text_run(p1, "................................", size=14)

        p2 = right_cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.line_spacing = 1.0
        self._add_text_run(
            p2,
            data.get("footer_cdo_name", data.get("cdo_name", "")),
            size=14
        )

        p3 = right_cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after = Pt(0)
        p3.paragraph_format.line_spacing = 1.0
        self._add_text_run(p3, "प्रमुख जिल्ला अधिकारी", size=14)

        date_info = {
            "year": data.get("document_date_year", ""),
            "month": data.get("document_date_month", ""),
            "day": data.get("document_date_day", ""),
            "day_num": data.get("document_date_day_num", "")
        }

        month_raw = date_info["month"]
        month_name = ""

        try:
            month_index = int(str(month_raw).strip())
            if 1 <= month_index <= len(NEPALI_MONTHS):
                month_name = NEPALI_MONTHS[month_index - 1]
            else:
                month_name = str(month_raw).strip()
        except (ValueError, TypeError):
            month_name = str(month_raw).strip() if month_raw else ""

        year_np = self._to_nepali_number(date_info["year"]) if date_info["year"] != "" else ""
        day_np = self._to_nepali_number(date_info["day"]) if date_info["day"] != "" else ""
        daynum_np = self._to_nepali_number(date_info["day_num"]) if date_info["day_num"] != "" else ""

        date_text = (
            f"ईति सम्वत {year_np} साल {month_name} "
            f"{day_np} गते रोज {daynum_np} शुभम्"
        )

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(0)
        self._add_text_run(p, date_text, size=14)

    def _to_nepali_number(self, num: int) -> str:
        """Convert integer to Nepali number string."""
        from config.constants import NEPALI_NUMBERS
        return "".join(NEPALI_NUMBERS.get(d, d) for d in str(num))